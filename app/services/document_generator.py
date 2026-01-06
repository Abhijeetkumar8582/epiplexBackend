from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import List, Dict
import os

class DocumentGenerator:
    def __init__(self):
        pass
    
    def generate_document(
        self,
        job_id: str,
        transcript: str,
        frame_analyses: List[Dict],
        output_dir: Path
    ) -> Dict[str, str]:
        """Generate documents in multiple formats"""
        output_paths = {}
        
        # Generate DOCX
        docx_path = self._generate_docx(job_id, transcript, frame_analyses, output_dir)
        output_paths["docx"] = str(docx_path)
        
        # Generate HTML
        html_path = self._generate_html(job_id, transcript, frame_analyses, output_dir)
        output_paths["html"] = str(html_path)
        
        return output_paths
    
    def _generate_docx(
        self,
        job_id: str,
        transcript: str,
        frame_analyses: List[Dict],
        output_dir: Path
    ) -> Path:
        """Generate DOCX document"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Video Processing Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Transcript section
        doc.add_heading('Transcript', level=1)
        doc.add_paragraph(transcript)
        doc.add_paragraph()
        
        # Frame analyses section
        doc.add_heading('Frame-by-Frame Analysis', level=1)
        
        for analysis in frame_analyses:
            # Timestamp heading
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f"Time: {analysis['timestamp']}")
            timestamp_run.bold = True
            timestamp_run.font.size = Pt(12)
            
            # Description
            doc.add_paragraph(analysis['description'])
            doc.add_paragraph()
        
        # Save document
        output_path = output_dir / f"{job_id}.docx"
        doc.save(str(output_path))
        return output_path
    
    def _generate_html(
        self,
        job_id: str,
        transcript: str,
        frame_analyses: List[Dict],
        output_dir: Path
    ) -> Path:
        """Generate HTML document"""
        output_path = output_dir / f"{job_id}.html"
        
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Video Processing Report</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1 {{
            text-align: center;
            color: #333;
            border-bottom: 3px solid #4CAF50;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #4CAF50;
            margin-top: 30px;
        }}
        .timestamp {{
            font-weight: bold;
            color: #2196F3;
            margin-top: 20px;
        }}
        .description {{
            margin-left: 20px;
            margin-bottom: 20px;
            padding: 10px;
            background-color: #f5f5f5;
            border-left: 4px solid #4CAF50;
        }}
        .transcript {{
            background-color: #f9f9f9;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
    </style>
</head>
<body>
    <h1>Video Processing Report</h1>
    
    <h2>Transcript</h2>
    <div class="transcript">
        <p>{transcript.replace(chr(10), '<br>')}</p>
    </div>
    
    <h2>Frame-by-Frame Analysis</h2>
"""
        
        for analysis in frame_analyses:
            html_content += f"""
    <div class="timestamp">Time: {analysis['timestamp']}</div>
    <div class="description">
        <p>{analysis['description'].replace(chr(10), '<br>')}</p>
    </div>
"""
        
        html_content += """
</body>
</html>
"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path

